from __future__ import unicode_literals
import codecs, sys, os, random, re, gzip, xlrd, math
import docx
import pickle
from collections import Counter
import subprocess
import sys
from text import Token, Sentence, Paragraph, Text




pipeline_dir            = '/Users/rex/Desktop/udpipe-master/src/udpipe'
udpipe_dir              = pipeline_dir
model_dir               = '/Users/rex/Desktop/AES_rex/scripts/sv_udpipe'
METADATA_TEMP_FORMAT    = "\nMETAMETADATADATA\n"
PARAGRAPH_TEMP_FORMAT   = "PARAGRAPH"
efselab_dir             = '/Users/rex/Desktop/AES_rex/scripts/efselab/swe_pipeline.py'
histnorm                = '/Users/rex/Desktop/AES_rex/scripts/HistNorm/'
output_dir              = '/Users/rex/Desktop/AES_rex/scripts/'
robert_infile           = '/Users/rex/Desktop/AES_rex/scripts/txt/'
robert_outfile          = '/Users/rex/Desktop/AES_rex/scripts/conll/R'


def text_anne(dir=None): #extract all text in directory dir
    #dir='Users/rex/Desktop/AES_rex/anne_data/'
    if dir:
        base = dir
    else:
        base = os.path.realpath(__file__).replace(os.path.basename(os.path.realpath(__file__)), '')
    dirs, paths = [base], []
    while True:
        new_dirs = []
        for dir in dirs:
            path_list = os.listdir(dir)
            for path in path_list:
                real_path = '/'.join([dir,path])
                if os.path.isdir(real_path):
                    new_dirs.append(real_path)
                elif os.path.isfile(real_path) and path != '.DS_Store':
                    paths.append(real_path)
        if new_dirs:
            dirs = new_dirs
        else:
            break
    texts = []
    for path in paths:
        print(path)
        if "SFI D" in path:
            #word2texts(path,'data', SFI=True)
            continue
        texts.extend(word2texts(path))
    return texts



def word2texts(filename):
    doc = docx.Document(filename)
    text = ''
    text_attr = ''
    title = False
    text_list = []
    if doc.paragraphs[0].text.startswith('<Text-id'):
        doc.paragraphs.pop(0)
    for para in doc.paragraphs:
        #print(para.text)
        if re.search(r'<Text-id.+>',para.text):
            continue
        if re.fullmatch(r'^ *<.+> *$', para.text):
            if text_attr:

                t = Text(text_attr[0], text_attr[1:])
                t.paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                #with open(outfile+'/'+text_id, 'w') as f:
                #    f.write(text)
                text_list.append(t)
                text = ''
            #text += re.search(r'<.+>', para.text).string
            start,end = re.search(r'<.+>', para.text).span()

            text_attr = [attr for attr in para.text[start+1:end-1].split(' ') if attr]
            text += '\n\n'
            title = True

        elif para.text.strip() and title:
            sents = para.text.strip().split('\n')
            sents = [sent.strip() for sent in sents]
            text += sents[0]
            text += '\n\n'
            if sents[1:]:
                text += '\n'.join(sents[1:])
                text += '\n\n'
            title = False

        elif para.text.strip():
            text += para.text.strip()
            text += '\n\n'

    #complete the last essay in the doc
    if text:
        t = Text(text_attr[0], text_attr[1:])
        t.paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        text_list.append(t)
    return text_list


class Text:

    def __init__(self, id, attributes):
        self.id = id
        sample, time, genre, grade, sex, subject, permission, place, education, format = attributes
        self.sample = sample
        self.time = time
        self.genre = genre
        self.grade = grade
        self.sex = sex
        self.subject = subject
        self.permission = permission
        self.place = place
        self.education = education
        self.format = format
        self.title = None
        self.paragraphs = []

    def __str__(self):
        return '\n'.join(self.paragraphs)


def text_filter(filename): #filter out the duplicate texts with same ids
    texts = text_anne(filename)
    id_texts = {}
    for text in texts:
        if text.id in id_texts:
            id_texts[text.id].append(text)
        else:
            id_texts[text.id] = [text]
    ids = [text.id for text in texts]
    duplicate_ids = [id for id, count in Counter(ids).most_common() if count > 1]

    for id in duplicate_ids:
        contain_texts = id_texts[id]
        id_texts[id] = [contain_texts.pop()]
        while contain_texts:
            t = contain_texts.pop()
            t_new = True
            for t_old in id_texts[id]:
                if t1ist2(t_old, t):
                    t_new = False
                    break
            if t_new:
                id_texts[id].append(t)

    #remove the duplicate texts with same id and update texts and ids
    texts = [text for texts in id_texts.values() for text in texts]
    ids = [text.id for text in texts]
    #create new ids for a different text with same id

    for text_id, freq in [(text_id, freq) for text_id, freq in Counter(ids).most_common() if freq > 1]:

        base_id = max([len(id[1:]) for id in ids if id[0] == text_id[0]])
        for i in range(1,freq):
            text = id_texts[text_id].pop()
            new_id = text_id[0] + str(base_id+i)
            id_texts[new_id] = [text]

    return [text for texts in id_texts.values() for text in texts]



def read_text(s): #input is a list
    text, paragraph, sentence = [], [], []
    for token in s:
        print(token.strip())
        if not token.strip() and not sentence:
            continue
        if not token.strip() and sentence:
            paragraph.append(sentence)
            sentence = []
            continue

        try:
            form = token.strip().split('\t')[1]
            if form == 'PARAGRAPH':
                if paragraph:
                    text.append(paragraph)
                    paragraph, sentence = [], []
            else:
                sentence.append(token.strip().split('\t'))
        except IndexError:
            print('Error')
            return False
      
    if sentence:
        paragraph.append(sentence)
    if paragraph:
        text.append(paragraph)
    return text


def text_extraction(filename=None, text_id=None):

    text = read_text(file_to_list(filename))
    if text is False:
        return False
    paras = []
    for para in text:
        sents = []
        for sent in para:
            words = []
            for word in sent:
                w = Token(word)
                words.append(w)
            s = Sentence(words)
            sents.append(s)
        p = Paragraph(sents)
        paras.append(p)
    print(paras)
    return Text(text_id, paras)


def data_set(data_dir, startwith=None, out='.'):
    '''This is a main function for creating pickle files for each essay'''
    data = text_filter(data_dir)
    if startwith:
        data = [text for text in data if text.id.startswith(startwith)]

    #remove texts without grades
    processed = os.listdir(out)
    print(len(data), len(processed))
    data = [text for text in data if text.grade != '_']

    texts = []
    print('Processing begins')
    for text in data: #
        print('%d texts left' %(len(data)-len(texts)))

        list_to_file('\n\n'.join(text.paragraphs), 'cache.txt') #used for a text object
        list_to_file(set_paragraphs('cache.txt'),'cache.txt')   #set boundaries for paragraphs
        #list_to_file(file_to_list(num_error+'/'+text.id),'cache.txt')
        if find_para_error(text):                               #check if there is wrong para tokenization
            con_broken_sents('cache.txt')
        pipeline(filename='cache.txt')
        t = text_extraction(filename='cache.conll', text_id=text.id)
        #####insert attributes for the text objects
        if t is False:
            subprocess.call(['mv','cache.txt',text.id])
            subprocess.call(['mv',text.id, out+'/'+'error'])
            continue
        t.sample = text.sample
        t.grade = text.grade
        t.time = text.time
        t.genre = text.genre
        t.education = text.education
        t.format = text.education
        t.sex = text.sex
        t.subject = text.subject
        t.permission = text.permission
        t.place = text.place
        t.paragraphs = text.paragraphs
        texts.append(check_ud_trees(t))
        with open(out+'/'+t.id, 'wb') as f:
            pickle.dump(t,f)


    return texts
##################################################################

class SFIText:

    def __init__(self, id, attributes):
        self.id = id
        test, track, source, genre, sex, lang, birth, school, grade = attributes
        self.test = test
        self.track = track
        self.source = source
        self.genre = genre
        self.sex = sex
        self.lang = lang
        self.birth = birth
        self.school = school
        self.grade = grade
        self.paragraphs = []

    def __str__(self):
        return "\n".join(self.paragraphs)







##########################################
##########pipeline########################
##########################################
'''create data by tokenization, tagging and parsing'''
def spellcheck(tokenized_file, spellchecked_file, HISTNORM_DIR):
    with open(spellchecked_file, 'w') as f:
        subprocess.call(['perl', '-CSAD', HISTNORM_DIR + 'scripts/normalise_levenshtein_elevtexter.perl',
                     tokenized_file, HISTNORM_DIR + 'resources/swedish/levenshtein/swedish.dic',
                     HISTNORM_DIR + 'resources/swedish/levenshtein/swedish.train.txt',
                     HISTNORM_DIR + 'resources/swedish/levenshtein/swedish.corp', 'noweights',
                     HISTNORM_DIR + 'resources/swedish/levenshtein/threshold.swedish.elevtexter.txt'],
                     stdout=f)
    # check cache.spell has the same length as cache.tok

    spe_text = file_to_list(spellchecked_file)
    tok_text = file_to_list(tokenized_file)
    if len(spe_text) != len(tok_text):
        tok, spe = alignword(tok_text, spe_text)
        list_to_file(spe,'cache.spell')
        list_to_file(tok,'cache.tok')


def pipeline(filename):
    '''perform tokenziation, tagging and parsing
    texts are text names with attribute of id
    udpipe_dir is path to the module for tokenizer, parser and pos tagger
    output_dir is a dir path to store the output'''
    #input: in_text; output: text.tok
    subprocess.call(['python3', efselab_dir, '--tokenized','-o', output_dir, filename])

    #normalize the tokenized words      input:cache.tok output:cache.spell
    spellcheck(output_dir+'cache.tok',output_dir+'cache.spell', histnorm)


    #tagging and parsing
    #input: cache.norm output: filename.conllu
    subprocess.call(['python3',efselab_dir,'--tagged', '--lemmatized', '--parsed', '--skip-tokenization','-o', output_dir, output_dir+'cache.spell'])

    #insert form into conll file
    restore_original_tokens('cache.conll','cache.tok')









###########################################
####function helper########################
###########################################
def t1ist2(t1,t2,robert1=False,robert2=False, showValue=False): #t1 and t2 are two objects here
    '''To compare if two texts are very similar by looking into
    the vocabulary and the frequency of each word.
    assumption: cover rate between t1 and t2 in terms of vocabulary 95%;
    the frequency of each word in the text'''

    def get_vocab(text, withFrequence=False, robert=False):
        vocab = [] # exclude punctuation
        if robert:
            vocab = re.split('\W', text.text)
        else:
            for para in text.paragraphs:
                words = re.split('\W', para)
                vocab.extend(words)
        if withFrequence:
            return {word:freq for word, freq in Counter(vocab).most_common()}
        return set(vocab)
    def vocab_cover(t1,t2):
        v1, v2 = get_vocab(t1,robert=robert1), get_vocab(t2,robert=robert2)
        try:
            if not v1.union(v2):
                print('The union set of text %s and %s on vocabulary is empty.' % (t1.id, t2.id))
                raise ValueError
        except ValueError:
            return 1
        return round(len(v1.intersection(v2)) / len(v1.union(v2)), 2)

    def word_freq(t1,t2):
        v1, v2 = get_vocab(t1, withFrequence=True, robert=robert1), get_vocab(t2, withFrequence=True, robert=robert2)
        #select the least around 50 procent words in both texts
        def select_word(v):
            sorted_l = [key for key, value in sorted(v.items(),key=lambda x:x[1])]
            length = len(sorted_l)
            return sorted_l[int(length/6):int(length/2)]

        words = set(select_word(v1)+(select_word(v2)))
        if not words:
            return 1
        def word_freq_similiarity(word,v1,v2):
            c1, c2 = v1.setdefault(word,0), v2.setdefault(word,0)
            if min([c1,c2]) == 0:
                return 0
            ratio = max([c1,c2]) / min([c1,c2])
            try:
                sim = round(2/(math.exp(2*(ratio-1))+1), 2)
            except OverflowError:
                print('OverflowError')
                return 0
            return sim
        similarities = [word_freq_similiarity(word, v1, v2) for word in words]
        return round(sum(similarities)/len(similarities), 2)

    if showValue:
        return vocab_cover(t1,t2) * 0.5 + word_freq(t1, t2) * 0.5
    return (vocab_cover(t1,t2) * 0.5 + word_freq(t1,t2) * 0.5) > 0.8




def is_a_ud_tree(heads):
    if not heads:
        return False
    children = list(range(1,len(heads)+1))

    if 0 not in heads:
        print("Root is missing.")
        return False
    if Counter(heads)[0] > 1:
        print("More than one root.")
        print(heads)
        return False
    if len(children) > len(set(children)):
        print("same indeces for two nodes.")
        return False
    for i in children:
        head = [heads[i - 1]]
        while 0 not in head:
            if heads[head[-1] - 1] in head:
                print("Cycle Error")
                return False
            head.append(heads[head[-1] - 1])
        head = []
    return True


def check_ud_trees(text):
    # check if all texts have grades

    for s in text.sentences:
        if is_a_ud_tree([int(token.head) for token in s.tokens]):
            continue
        else:
            text.ud = True
            s.ud = True
            print('\t'.join([token.norm for token in s.tokens]))
            print('\t'.join([token.head for token in s.tokens]))

            # use udpipe to reparse the sentence
            with open('cache.udpipe', 'w') as f:
                f.write('\n'.join([word.norm for word in s.tokens]))
            with open('cache.udpipe.conll', 'w') as f:
                subprocess.call([udpipe_dir, '--tag', '--parse', '--input', 'vertical', model_dir, 'cache.udpipe'],
                                stdout=f)
            with open('cache.udpipe.conll', 'r') as f:
                lines = f.readlines()
            heads = [int(line.split('\t')[6]) for line in lines if line.strip() if line[:2] != '# ']
            print('\t'.join([str(head) for head in heads]))
            for head, token in zip(heads, s.tokens):
                token.head = head

    return text



def find_para_error(text):
    '''functions to find texts containing broken sentences due to wrong sentence tokenization'''
    hasError = 0
    for para in text.paragraphs:
        if para[-1] not in '?.!>-]")…' and para[-1] not in "'":
            hasError += 1
        if hasError > 5:
            return True
    return False

def con_broken_sents(text):
    '''functions to concatenate two incorrectly separately paragraphs'''
    lines = file_to_list(text)
    lines = [line.strip() for line in lines if line.strip() if line.strip() != 'PARAGRAPH']
    content = lines[:2]
    lines = lines[2:]
    while lines:
        line = lines.pop(0)
        if content[-1][-1] not in '!.?' or re.search(r'[a-zäöå]', line[0]):
            #concatenate
            content[-1] = ' '.join([content[-1],line])
        else:
            content.append(line)
    list_to_file('\n\nPARAGRAPH\n\n'.join(content), text)





def list_to_file(list, file):
    f = open(file, 'w')
    f.write(''.join(map(lambda x: str(x), list)))
    f.close()

def file_to_list(file):
    f = open(file)
    list = f.readlines()
    return list

def set_paragraphs(input_file):
    infile = file_to_list(input_file)
    for x in range(len(infile)):
        if infile[x] == "\n":
            del infile[x]
            infile.insert(x, "\n" + PARAGRAPH_TEMP_FORMAT + "\n\n")
    return infile

def restore_paragraphs(input_file):
    decrement = 0
    infile = file_to_list(input_file)
    for x in range(len(infile)):
        x -= decrement

        if PARAGRAPH_TEMP_FORMAT in [l.strip() for l in infile[x].split("\t")]:

            del infile[x]
            decrement += 1
    return infile


def levenshtein(t1, t2):
    if not t1 or not t2: return max([len(t1),len(t2)])
    row, col = len(t1)+1, len(t2)+1
    matrix = [[0] * col for _ in range(row)]
    matrix[0] = [i for i in range(col)]
    for i in range(row):
        matrix[i][0] = i
    for r in range(1,row):
        for c in range(1,col):
            if t1[r-1] == t2[c-1]:
                matrix[r][c] = min([matrix[r-1][c-1],
                                    matrix[r][c-1]+1,
                                    matrix[r-1][c]+1])
            else:
                matrix[r][c] = min([matrix[r-1][c-1]+1,
                                    matrix[r][c-1]+1,
                                    matrix[r-1][c]+1])
    return matrix[-1][-1]



def alignword(tok, spell):
    row, col = len(tok)+1, len(spell)+1
    matrix = [[(0,0,0)]*col for _ in range(row)]
    matrix[0] = [(0,0,0)]+[(i*2,0,i-1) for i in range(1,col)]
    for i in range(1,row):
        matrix[i][0] = i*2,i-1,0
    for r in range(1,row):
        for c in range(1, col):
            if r==c==1:
                matrix[r][c] = 0,0,0
                continue
            d = levenshtein(tok[r-1],spell[c-1]) + matrix[r-1][c-1][0]
            dr = matrix[r][c-1][0]+2
            dc = matrix[r-1][c][0]+2
            if d == min([d,dr,dc]):
                matrix[r][c] = d,r-1, c-1
            elif dr == min([d,dr,dc]):
                matrix[r][c] = dr, r, c-1
            else:
                matrix[r][c] = dc, r-1, c

    index_pair = []
    r,c = matrix[-1][-1][1:3]
    while min([r,c]) != 0:
        index_pair.append((r,c))
        r,c = matrix[r][c][1:3]
    index_pair.reverse()
    tokens, spells = [], []
    if (0,0) not in index_pair:
        index_pair.insert(0,(0,0))
    for t, s in index_pair:
        if tok[t] != spell[s] and re.fullmatch(r'[0-9|.|(|)|/|-| ]+', tok[t].strip()):
            spells.append(tok[t])
        else:
            spells.append(spell[s])
        tokens.append(tok[t])
    return tokens, spells


def restore_original_tokens(tagged_file, original_token_file):
    updated_tagged_file = []
    tagged = file_to_list(tagged_file)
    original = file_to_list(original_token_file)

    for tag, token in zip(tagged, original):
        if token.strip():
            tag_line = tag.split('\t')
            tag_line.insert(1, token.strip())
            updated_tagged_file.append('\t'.join(tag_line))
        else:
            updated_tagged_file.append(tag)

    list_to_file(updated_tagged_file, tagged_file)


def read_pickle(filename):
    with open(filename, 'rb') as f:
        return pickle.load(f)






def grade_distribution(texts, level=False):
    if level:
        texts = level_texts(texts,level)
    return Counter([text.grade for text in texts]).most_common()


def filter_same_level(texts, level=None, robert1=False, robert2=False):
    if level:
        level_texts = [text for text in texts if text.id[0] == level]
    else:
        level_texts = texts
    n_level_texts = len(level_texts)
    print('Level %s contains %d texts before level filtering' % (level, len(level_texts)))
    filtered_texts = [level_texts.pop()]
    while level_texts:
        current_text = level_texts.pop()
        print('%d texts are left. %s' %(len(level_texts), current_text.id))
        current_new = True
        for target in filtered_texts:
            if t1ist2(current_text, target, robert1, robert2):
                current_new = False
                print('Text %s and text %s have a high similarity!' %(target.id, current_text.id))
        if current_new:
            filtered_texts.append(target)
    print('Level %s contains %d texts that are filtered' % (level, n_level_texts-len(filtered_texts)))
    return filtered_texts





######################################################################
######################################################################
######################################################################
######################################################################


#if __name__ == '__main__':
#    main()
